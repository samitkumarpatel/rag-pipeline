"""Unit tests for DocxParser."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.parsers.docx_parser import DocxParser


@pytest.fixture
def parser() -> DocxParser:
    return DocxParser()


@pytest.fixture
def simple_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the first paragraph of the document.")
    doc.add_paragraph("This is the second paragraph with more content.")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_table(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Document with a table.")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    table.cell(0, 2).text = "City"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "30"
    table.cell(1, 2).text = "Berlin"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    return path


def test_supports_docx_mime(parser: DocxParser) -> None:
    assert parser.supports(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ) is True


def test_supports_msword_mime(parser: DocxParser) -> None:
    assert parser.supports("application/msword") is True


def test_does_not_support_pdf(parser: DocxParser) -> None:
    assert parser.supports("application/pdf") is False


def test_raises_file_not_found(parser: DocxParser, tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        parser.parse(tmp_path / "nonexistent.docx")


def test_parse_simple_docx(parser: DocxParser, simple_docx: Path) -> None:
    result = parser.parse(simple_docx)

    assert result.page_count == 1
    assert "Introduction" in result.full_text
    assert "first paragraph" in result.full_text
    assert "second paragraph" in result.full_text


def test_parse_docx_extracts_table_content(
    parser: DocxParser, docx_with_table: Path
) -> None:
    result = parser.parse(docx_with_table)

    assert "Alice" in result.full_text
    assert "Berlin" in result.full_text
    assert "Name" in result.full_text


def test_metadata_contains_paragraph_count(
    parser: DocxParser, simple_docx: Path
) -> None:
    result = parser.parse(simple_docx)
    assert "paragraph_count" in result.metadata
    assert int(result.metadata["paragraph_count"]) >= 2


def test_raises_on_corrupt_docx(parser: DocxParser, tmp_path: Path) -> None:
    bad_file = tmp_path / "corrupt.docx"
    bad_file.write_bytes(b"PK not a real docx at all %%corrupt%%")
    with pytest.raises(RuntimeError):
        parser.parse(bad_file)
