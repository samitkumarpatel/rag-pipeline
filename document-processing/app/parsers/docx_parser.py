"""
DOCX parser using python-docx.

Extracts:
- Paragraph text in document order
- Table cell content (row-major order)
- Heading levels for section metadata
"""
from __future__ import annotations

from pathlib import Path

import structlog
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.parsers.base import AbstractParser, ParsedDocument

logger = structlog.get_logger(__name__)

_SUPPORTED_MIME_TYPES = frozenset({
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
})


class DocxParser(AbstractParser):
    """Parse DOCX / DOC documents into structured text."""

    def supports(self, mime_type: str) -> bool:
        return mime_type in _SUPPORTED_MIME_TYPES

    def parse(self, file_path: Path) -> ParsedDocument:
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")

        logger.info("parsing_docx", path=str(file_path))

        try:
            doc = Document(str(file_path))
        except PackageNotFoundError as exc:
            raise RuntimeError(f"Corrupt or unreadable DOCX: {file_path}") from exc
        except Exception as exc:
            raise RuntimeError(f"Failed to open DOCX: {file_path}: {exc}") from exc

        # DOCX has no intrinsic page numbers — emit everything as a single page
        # and preserve heading structure for chunk metadata
        blocks: list[str] = []
        section_title: str | None = None
        metadata: dict[str, object] = {}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Capture heading text as section context
            if para.style and para.style.name.startswith("Heading"):
                section_title = text
                blocks.append(f"\n## {text}\n")
            else:
                blocks.append(text)

        # Extract table content
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    table_rows.append(row_text)
            if table_rows:
                blocks.append("\n".join(table_rows))

        full_text = "\n".join(blocks)
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)

        parsed = ParsedDocument(pages=[(1, full_text)], metadata=metadata)
        logger.info(
            "docx_parsed",
            path=str(file_path),
            total_chars=len(full_text),
            paragraphs=metadata["paragraph_count"],
        )
        return parsed
